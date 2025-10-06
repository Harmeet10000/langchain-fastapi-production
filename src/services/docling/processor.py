"""Document processing service using Docling and other parsers."""

from typing import List, Dict, Any, Optional
from pathlib import Path
import tempfile
import hashlib
from io import BytesIO

from langchain.text_splitter import RecursiveCharacterTextSplitter, CharacterTextSplitter
from langchain.schema import Document
import pypdf
from docx import Document as DocxDocument
import pandas as pd
from bs4 import BeautifulSoup
import markdown

from src.core.config.logging_config import LoggerAdapter
from src.core.cache.redis_client import redis_cache

logger = LoggerAdapter(__name__)


class DocumentProcessor:
    """Service for processing various document formats."""

    def __init__(self):
        """Initialize document processor."""
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )

    async def process_document(
        self,
        file_path: str,
        file_type: Optional[str] = None,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        use_cache: bool = True
    ) -> List[Dict[str, Any]]:
        """Process document and extract content with metadata."""
        try:
            # Get file hash for caching
            file_hash = self._get_file_hash(file_path)
            cache_key = f"doc:processed:{file_hash}"

            # Check cache
            if use_cache:
                cached_result = await redis_cache.get(cache_key)
                if cached_result:
                    logger.info("Returning cached processed document")
                    return cached_result

            # Detect file type if not provided
            if not file_type:
                file_type = Path(file_path).suffix.lower()

            # Process based on file type
            if file_type in [".pdf", "pdf"]:
                content, metadata = await self._process_pdf(file_path)
            elif file_type in [".docx", ".doc", "docx", "doc"]:
                content, metadata = await self._process_docx(file_path)
            elif file_type in [".txt", "txt"]:
                content, metadata = await self._process_text(file_path)
            elif file_type in [".md", "md", "markdown"]:
                content, metadata = await self._process_markdown(file_path)
            elif file_type in [".html", "html"]:
                content, metadata = await self._process_html(file_path)
            elif file_type in [".xlsx", ".xls", "xlsx", "xls"]:
                content, metadata = await self._process_excel(file_path)
            elif file_type in [".csv", "csv"]:
                content, metadata = await self._process_csv(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")

            # Create custom text splitter if needed
            if chunk_size != 1000 or chunk_overlap != 200:
                text_splitter = RecursiveCharacterTextSplitter(
                    chunk_size=chunk_size,
                    chunk_overlap=chunk_overlap,
                    length_function=len
                )
            else:
                text_splitter = self.text_splitter

            # Split content into chunks
            chunks = text_splitter.split_text(content)

            # Create documents with metadata
            documents = []
            for i, chunk in enumerate(chunks):
                doc_metadata = {
                    **metadata,
                    "chunk_index": i,
                    "chunk_total": len(chunks),
                    "chunk_size": len(chunk)
                }

                documents.append({
                    "content": chunk,
                    "metadata": doc_metadata
                })

            # Cache result
            if use_cache:
                await redis_cache.set(cache_key, documents, ttl=7200)  # 2 hours

            logger.info(f"Processed document: {metadata['filename']}, {len(chunks)} chunks")
            return documents

        except Exception as e:
            logger.error("Failed to process document", error=str(e))
            raise

    async def _process_pdf(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Process PDF document."""
        try:
            content_parts = []
            metadata = {
                "filename": Path(file_path).name,
                "file_type": "pdf",
                "page_count": 0
            }

            with open(file_path, "rb") as file:
                pdf_reader = pypdf.PdfReader(file)
                metadata["page_count"] = len(pdf_reader.pages)

                for page_num, page in enumerate(pdf_reader.pages, 1):
                    text = page.extract_text()
                    if text:
                        content_parts.append(f"[Page {page_num}]\n{text}")

            content = "\n\n".join(content_parts)
            return content, metadata

        except Exception as e:
            logger.error("Failed to process PDF", error=str(e))
            raise

    async def _process_docx(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Process Word document."""
        try:
            doc = DocxDocument(file_path)
            content_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content_parts.append(paragraph.text)

            # Extract tables
            for table in doc.tables:
                table_text = self._extract_table_text(table)
                if table_text:
                    content_parts.append(table_text)

            content = "\n\n".join(content_parts)
            metadata = {
                "filename": Path(file_path).name,
                "file_type": "docx",
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables)
            }

            return content, metadata

        except Exception as e:
            logger.error("Failed to process DOCX", error=str(e))
            raise

    async def _process_text(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Process text file."""
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                content = file.read()

            metadata = {
                "filename": Path(file_path).name,
                "file_type": "txt",
                "character_count": len(content),
                "line_count": content.count("\n") + 1
            }

            return content, metadata

        except Exception as e:
            logger.error("Failed to process text file", error=str(e))
            raise

    async def _process_markdown(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Process Markdown file."""
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                md_content = file.read()

            # Convert to HTML then extract text
            html = markdown.markdown(md_content)
            soup = BeautifulSoup(html, "html.parser")
            content = soup.get_text()

            metadata = {
                "filename": Path(file_path).name,
                "file_type": "markdown",
                "character_count": len(content),
                "has_code_blocks": "```" in md_content
            }

            return content, metadata

        except Exception as e:
            logger.error("Failed to process Markdown", error=str(e))
            raise

    async def _process_html(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Process HTML file."""
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                html_content = file.read()

            soup = BeautifulSoup(html_content, "html.parser")

            # Remove script and style elements
            for element in soup(["script", "style"]):
                element.decompose()

            # Get text
            content = soup.get_text()

            # Clean up whitespace
            lines = (line.strip() for line in content.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            content = "\n".join(chunk for chunk in chunks if chunk)

            metadata = {
                "filename": Path(file_path).name,
                "file_type": "html",
                "title": soup.title.string if soup.title else None,
                "character_count": len(content)
            }

            return content, metadata

        except Exception as e:
            logger.error("Failed to process HTML", error=str(e))
            raise

    async def _process_excel(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Process Excel file."""
        try:
            df_dict = pd.read_excel(file_path, sheet_name=None)
            content_parts = []

            for sheet_name, df in df_dict.items():
                content_parts.append(f"Sheet: {sheet_name}")
                content_parts.append(df.to_string())

            content = "\n\n".join(content_parts)
            metadata = {
                "filename": Path(file_path).name,
                "file_type": "excel",
                "sheet_count": len(df_dict),
                "total_rows": sum(len(df) for df in df_dict.values())
            }

            return content, metadata

        except Exception as e:
            logger.error("Failed to process Excel", error=str(e))
            raise

    async def _process_csv(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Process CSV file."""
        try:
            df = pd.read_csv(file_path)
            content = df.to_string()

            metadata = {
                "filename": Path(file_path).name,
                "file_type": "csv",
                "rows": len(df),
                "columns": len(df.columns),
                "column_names": list(df.columns)
            }

            return content, metadata

        except Exception as e:
            logger.error("Failed to process CSV", error=str(e))
            raise

    def _extract_table_text(self, table) -> str:
        """Extract text from a Word document table."""
        rows = []
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(row_text))
        return "\n".join(rows)

    def _get_file_hash(self, file_path: str) -> str:
        """Get hash of file for caching."""
        hasher = hashlib.sha256()
        with open(file_path, "rb") as file:
            for chunk in iter(lambda: file.read(4096), b""):
                hasher.update(chunk)
        return hasher.hexdigest()

    async def extract_metadata_only(self, file_path: str) -> Dict[str, Any]:
        """Extract only metadata without processing content."""
        try:
            file_type = Path(file_path).suffix.lower()
            metadata = {
                "filename": Path(file_path).name,
                "file_size": Path(file_path).stat().st_size,
                "file_type": file_type
            }

            if file_type == ".pdf":
                with open(file_path, "rb") as file:
                    pdf_reader = pypdf.PdfReader(file)
                    metadata["page_count"] = len(pdf_reader.pages)

            return metadata

        except Exception as e:
            logger.error("Failed to extract metadata", error=str(e))
            return {}


# Create global instance
document_processor = DocumentProcessor()
